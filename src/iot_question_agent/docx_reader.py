from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document


POINT_RE = re.compile(r"^\s*(\d+(?:\.\d+)+)\.?\s+(.+)$")
SECTION_RE = re.compile(r"^\s*(\d+)\.?\s+(.+)$")


@dataclass
class DocBlock:
    section: str
    point: str
    text: str


def _clean_text(text: str) -> str:
    return " ".join((text or "").replace("\xa0", " ").split()).strip()


def _is_heading(paragraph) -> bool:
    style_name = ""
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        return False
    lowered = style_name.lower()
    return "heading" in lowered or "заголов" in lowered


def read_docx_blocks(path: Path) -> list[DocBlock]:
    doc = Document(str(path))
    blocks: list[DocBlock] = []

    current_section = "Без раздела"
    current_point = "без номера"
    current_text_parts: list[str] = []

    def flush() -> None:
        nonlocal current_text_parts
        text = _clean_text(" ".join(current_text_parts))
        if text:
            blocks.append(DocBlock(current_section, current_point, text))
        current_text_parts = []

    for p in doc.paragraphs:
        text = _clean_text(p.text)
        if not text:
            continue

        point_match = POINT_RE.match(text)
        section_match = SECTION_RE.match(text)

        if point_match:
            flush()
            current_point = point_match.group(1).strip().rstrip(".")
            current_text_parts = [point_match.group(2).strip()]
            continue

        if section_match and (_is_heading(p) or len(text) < 160):
            flush()
            current_section = text
            current_point = "без номера"
            continue

        if _is_heading(p) and len(text) < 180:
            flush()
            current_section = text
            current_point = "без номера"
            continue

        current_text_parts.append(text)

    flush()

    # Таблицы в инструкциях встречаются редко, но если есть — добавляем их как текстовые блоки.
    table_index = 0
    for table in doc.tables:
        table_index += 1
        rows = []
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        table_text = _clean_text("; ".join(rows))
        if table_text:
            blocks.append(DocBlock(f"Таблица {table_index}", "без номера", table_text))

    return blocks


def blocks_to_text(blocks: list[DocBlock]) -> str:
    parts: list[str] = []
    for i, b in enumerate(blocks, start=1):
        parts.append(
            f"[Блок {i}]\nРаздел: {b.section}\nПункт: {b.point}\nТекст: {b.text}"
        )
    return "\n\n".join(parts)


def split_blocks_by_chars(blocks: list[DocBlock], max_chars: int) -> list[list[DocBlock]]:
    if max_chars < 1000:
        raise ValueError("max_chars_per_llm_call must be >= 1000")

    chunks: list[list[DocBlock]] = []
    current: list[DocBlock] = []
    current_len = 0

    for block in blocks:
        block_text = blocks_to_text([block])
        block_len = len(block_text)
        if current and current_len + block_len > max_chars:
            chunks.append(current)
            current = []
            current_len = 0
        current.append(block)
        current_len += block_len

    if current:
        chunks.append(current)
    return chunks
